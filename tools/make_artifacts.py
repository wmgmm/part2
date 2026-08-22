#!/usr/bin/env python3
"""Generate the seven venture artifact files for the workshop missions.

Run from the repo root:  python3 tools/make_artifacts.py
Outputs land in public/placeholders/. Deterministic (seeded), so re-running
produces identical files. All content is fictional teaching material.

To re-skin the workshop, edit the text/config below and re-run.
"""

import csv
import random
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors

OUT = Path(__file__).resolve().parent.parent / "public" / "placeholders"
OUT.mkdir(parents=True, exist_ok=True)
rng = random.Random(2026)

FOOTER = "Fictional teaching material for the AI in the Workplace workshop, Cardiff University."


# ---------------------------------------------------------------- docx helpers
def make_docx(path, title, blocks):
    doc = Document()
    doc.add_heading(title, level=0)
    for kind, text in blocks:
        if kind == "h":
            doc.add_heading(text, level=1)
        elif kind == "p":
            doc.add_paragraph(text)
        elif kind == "b":
            doc.add_paragraph(text, style="List Bullet")
    p = doc.add_paragraph()
    run = p.add_run(FOOTER)
    run.font.size = Pt(8)
    run.italic = True
    doc.save(path)


# ----------------------------------------------------------------- pdf helpers
def make_pdf(path, title, blocks):
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=10.5, leading=15)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], spaceBefore=10)
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=20 * mm, rightMargin=20 * mm,
                            topMargin=18 * mm, bottomMargin=18 * mm,
                            title=title)
    story = [Paragraph(title, styles["Title"]), Spacer(1, 6)]
    for kind, text in blocks:
        if kind == "h":
            story.append(Paragraph(text, h1))
        elif kind == "p":
            story.append(Paragraph(text, body))
            story.append(Spacer(1, 5))
        elif kind == "b":
            story.append(Paragraph("• " + text, body))
        elif kind == "t":  # table: list of rows
            tbl = Table(text, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 6))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"<i>{FOOTER}</i>", ParagraphStyle("f", parent=body, fontSize=7.5)))
    doc.build(story)


# ================================================================ 1. THE BRIEF
# Square brackets mark the swap points; the venture default stays inside them
# so the prompt runs verbatim. Matches the prompt shown on the site (real
# newlines here correspond to \n escapes in the missions.js string).
BRIEF_TEXT = (
    "ROLE: You are a market research analyst preparing a briefing for a first-time food vendor.\n"
    "TASK: Research the viability of [a small Welsh cake stall on a UK university campus (Cardiff)].\n"
    "FORMAT: Use exactly these section headings:\n"
    "- Who Buys, and When: footfall patterns for [campus food stalls] across the day and the academic year.\n"
    "- Price Expectations: realistic prices for [a fresh Welsh cake and comparable snacks].\n"
    "- Competitors: [on and near a typical campus, including mobile and permanent options].\n"
    "- Succeed or Fail: what decides it for [small campus food ventures], including [the regulatory basics for UK street food].\n"
    "- Blind Spots: three opportunities or risks I am probably not thinking about.\n"
    "CONSTRAINTS: Cite every claim to its source. If a price or figure is not in your sources, write DATA UNAVAILABLE rather than estimating it."
)

# Drift guard: the same brief text lives in src/data/missions.js (the m1
# PromptBox). If someone edits one copy, fail loudly here rather than shipping
# a docx that disagrees with the site.
import re as _re

_missions = (Path(__file__).resolve().parent.parent / "src" / "data" / "missions.js").read_text(encoding="utf-8")
_m = _re.search(r"promptLabel: 'THE BRIEF \(EDIT FOR YOUR IDEA\)',\s*prompt:\s*\n?\s*'((?:[^'\\]|\\.)*)'", _missions)
_site_brief = _m.group(1).replace("\\'", "'").replace("\\n", "\n") if _m else None
if _site_brief != BRIEF_TEXT:
    raise SystemExit(
        "BRIEF_TEXT drift: tools/make_artifacts.py no longer matches the m1 brief "
        "prompt in src/data/missions.js. Update both copies to say the same thing."
    )

make_docx(OUT / "Venture_Research_Brief.docx", "Market Research Brief: The Gravitas Venture", [
    ("p", "Paste the brief below into Gemini Deep Research, or rewrite it for your own idea. "
          "Edit the research plan it proposes before you press start."),
    ("h", "The brief"),
    *[("p", _line) for _line in BRIEF_TEXT.split("\n")],
    ("h", "Rewriting it for your own idea"),
    ("p", "Keep the shape: what decision the research informs, the customer, the price question, "
          "the competition, the regulations, and an open invitation for risks you have not "
          "thought of. Swap the nouns."),
])

# ======================================================== 2. THE MARKET REPORT
make_pdf(OUT / "Venture_Market_Report.pdf", "Campus Welsh Cake Stall: Market Briefing", [
    ("p", "A research briefing on the viability of a small Welsh cake stall on a Cardiff "
          "university campus. Prepared as a grounded stand-in for a live Deep Research run."),
    ("h", "1. Who buys, and when"),
    ("p", "Campus food purchasing is sharply peaked. Observed weekday footfall past the proposed "
          "pitch (main walkway, outside the students' union) averages 2,400 people per day in "
          "term time, with 61 per cent of passing trade between 11:30 and 14:00. Morning trade "
          "is modest; afternoon trade tails off after 15:30. Weekend footfall falls below 400."),
    ("p", "The academic calendar dominates demand. Teaching runs roughly 22 weeks across two "
          "semesters; once vacations, exam quiet periods and bank holidays are counted, the "
          "campus is effectively dormant for around 14 weeks of the year. Any annual projection "
          "built on a term-time week misstates the year badly."),
    ("h", "2. Price expectations"),
    ("p", "Comparable campus snack pricing clusters tightly: fresh bakery items sell at "
          "£1.60 to £2.40 on campus, with a premium fresh-off-the-griddle item sustainable at "
          "£2.20 to £2.60. Student purchasers are strongly price-anchored below £3; a £3.20 "
          "single-item price point sits outside the observed tolerance for daily purchases and "
          "would position the stall as an occasional treat."),
    ("h", "3. Competition"),
    ("p", "Two direct competitors operate within a five-minute walk of the proposed pitch: "
          "The Cwtch Coffee Co. (permanent unit inside the students' union, sells packaged "
          "Welsh cakes at £1.80) and Mrs Mabli's Bakes (a mobile van trading Tuesdays and "
          "Thursdays, fresh griddle cakes at £2.20). Substitute competition includes two "
          "supermarkets within 400 metres selling six-packs at £1.35. Differentiation must "
          "rest on freshness and theatre, not on being the only option."),
    ("h", "4. What makes small campus food ventures succeed or fail"),
    ("b", "Fail: underestimating production ceilings at peak; a queue longer than eight minutes "
          "reliably loses the back half of the line."),
    ("b", "Fail: single-operator dependence; illness or equipment failure with no cover closes "
          "the stall and resets habit formation."),
    ("b", "Fail: ignoring vacation troughs in cash-flow planning."),
    ("b", "Succeed: pre-order or reservation mechanisms that smooth the lunch spike."),
    ("b", "Succeed: card-first payment; cash-only stalls lose an estimated fifth of impulse trade."),
    ("b", "Succeed: visible allergen labelling; dietary-option availability drives repeat visits."),
    ("h", "5. Regulatory basics (UK street food)"),
    ("b", "Register as a food business with the local authority at least 28 days before trading."),
    ("b", "Comply with Natasha's Law (PPDS labelling) for any item packed before sale; maintain "
          "an allergen matrix for made-to-order items."),
    ("b", "Food hygiene training for all operators; expect an EHO inspection and rating."),
    ("b", "A campus pitch requires a commercial licence from the university's estates division, "
          "typically with a weekly pitch fee and public liability insurance."),
    ("h", "6. Three things you are probably not thinking about"),
    ("b", "Weather exposure: an outdoor Cardiff pitch loses an estimated 30 trading days a year "
          "to rain unless covered."),
    ("b", "Gluten-free demand: coeliac-safe griddle work needs separated equipment, and campus "
          "audiences ask; the queue for the one gluten-free-friendly van on comparable campuses "
          "is disproportionate to its share of the menu."),
    ("b", "Data protection: any loyalty scheme collecting student details is subject to UK GDPR; "
          "sharing customer lists with third parties without explicit consent is unlawful."),
])

# ==================================================== 3. THE CUSTOMER FEEDBACK
THEMES = {
    "queue": [
        "Queued 20 minutes at lunch. Cakes great, queue not.",
        "The line was halfway to the library. Gave up.",
        "Waited so long my lecture started. Please sort the queue.",
        "15 minute wait for two cakes. Lovely cakes though.",
        "Queue moves so slowly at 12. One person cannot do it all.",
        "I time my visits for 10am now because lunch is hopeless.",
        "Queue chaos at lunchtime, twice this week I walked away.",
        "Please do something about the wait. 18 minutes today.",
        "The queue is the only bad thing about this stall.",
        "Half my lunch break gone standing in line.",
        "Why is there no way to order ahead? The queue is wild.",
        "Stood in the rain in a queue. Cardiff, innit. Still worth it. Just.",
        "Queue too long, went to the union shop instead.",
        "If I could preorder I would buy double the amount.",
        "The wait at peak is genuinely 15-20 minutes.",
        "Queue put my friends off, they went to Cwtch Coffee.",
        "Everything is great except the twenty minute wait.",
        "One griddle cannot feed a lunchtime campus, bless him.",
        "Waited ages, then card machine queue on top. Sort it.",
        "The 12:30 queue needs its own postcode.",
        "Love the cakes, hate the queue, simple as.",
        "Queue again. Third time. Still joined it. Send help.",
        "Would it kill you to have a fast lane for plain cakes?",
        "25 min at 12:15 today. That is a lecture, not a snack break.",
        "The queue is a bonding experience I did not ask for.",
        "Waiting time at lunch is silly now the word has spread.",
        "Great product, queue management non-existent.",
        "I now budget half an hour for a Welsh cake. Reflect on that.",
    ],
    "queue_price": [  # dual-theme: queue + price
        "£3 and I still waited fifteen minutes for a cold one.",
        "Pricey AND a queue? One or the other, please.",
        "Waited 20 minutes to pay more than the union shop charges.",
    ],
    "queue_soldout": [  # dual-theme: queue + sold out
        "Queued 10 minutes to be told the tray was empty.",
        "The queue at noon is really a lottery for the last dozen.",
        "Got to the front and the plain ones were gone. Again.",
    ],
    "price": [
        "£2.50 is fair for fresh off the griddle to be honest.",
        "Bit dear compared to the supermarket six pack.",
        "Cheaper than the coffee shop muffin and twice as nice.",
        "Price crept up this term. Noticed.",
        "Fine for a treat, too much for every day.",
        "Good value for the box of six, singles feel steep.",
        "£2.50 for one cake is a lot on a student budget.",
        "Worth every penny, fight me.",
        "Would buy more if there was a loyalty discount.",
        "Price is fine, size could be bigger.",
        "Honestly decent value for something made in front of you.",
        "Costs more than my bus fare home. Still bought three.",
        "The box of six is the only sensible way to buy these.",
        "A pound more than the packaged ones but ten times better.",
        "Too expensive for a daily habit, sadly.",
        "Fair price, fresh and warm. No complaints on cost.",
        "Students get a discount? No? Should we?",
        "I would pay more for a bigger one, just saying.",
        "Price fine. Portion fine. I just wanted to leave a comment.",
    ],
    "soldout": [
        "Sold out by 1pm AGAIN.",
        "Third visit this week where everything was gone by early afternoon.",
        "Why bother opening until 3 if the cakes are gone at 1?",
        "Turned up at 1:15, empty trays, sad walk to the library.",
        "Make more! You sell out every single day!",
        "The vanilla ones are gone by noon without fail.",
        "Sold out when I arrived. The sign said open. My heart said closed.",
        "Every time I have a late lunch, no cakes left.",
        "Gone by one o'clock. This is a capacity problem, not a marketing one.",
        "You could sell double the stock. You know it. We know it.",
        "Afternoon customers exist! Save us a tray!",
        "No gluten free left by 11:45, all gone by 1.",
        "Ran out early on both my visits.",
        "1pm: empty. 1:05pm: still empty. Checked twice.",
        "If you made twice as many I would still worry.",
        "The sell-out is the best marketing and the worst experience.",
    ],
    "dietary": [
        "More gluten free please, they go instantly.",
        "Is there a vegan option? Would buy weekly.",
        "Coeliac here. Is the gluten free one done on a separate griddle?",
        "Allergen labels are hard to find on the stall. Please display them.",
        "Nut allergy: nobody could tell me what was safe. Went without.",
        "GF cakes lovely but there were only six of them. Six!",
        "Dairy free option would be amazing.",
        "Please list ingredients somewhere visible.",
        "The gluten free one is great, stock more than a handful.",
        "Vegan Welsh cakes exist, I have seen them in the wild. Please.",
        "Asked about allergens and got a shrug. Not good enough.",
        "More dietary options and you have my entire loan.",
        "GF sold out before my seminar ended. Predictable by now.",
        "Half my flat cannot eat these. Diversify the menu?",
    ],
    "payment": [
        "Cash only?! In this economy?",
        "The card machine was down again today.",
        "Please get contactless sorted, nobody carries cash.",
        "Card minimum of £5 forces me to buy three cakes. Clever. Annoying.",
        "Payment took ages, the machine kept timing out.",
        "Apple Pay please!",
        "Had to walk to a cash point and back. The cakes were gone by then.",
        "Card machine queue is slower than the cake queue.",
        "Contactless failed, held everyone up, cash saved the day. In 2026!",
    ],
    "location": [
        "Move nearer the science buildings, we have money too.",
        "Hard to find unless you already know it is there.",
        "Opening at 10 misses the 9am crowd entirely.",
        "Could you do Fridays? The one day I am on campus!",
        "The pitch is windy, the napkins fly, the vibe suffers.",
        "Wish you opened earlier. My 9am needs this.",
        "A sign would help. I walked past it for a month.",
    ],
    "praise": [
        "Best thing on this campus, no notes.",
        "The smell alone is worth the walk.",
        "Reminds me of my nan's, which is the highest possible mark.",
        "Lovely staff, lovely cakes, lovely everything.",
        "10/10 would queue again. Unfortunately.",
    ],
}

# rating tendencies per theme
THEME_RATING = {
    "queue": [2, 2, 3, 3, 4], "queue_price": [2, 2, 3], "queue_soldout": [2, 2, 3],
    "price": [3, 3, 4, 4, 5], "soldout": [2, 2, 3, 3], "dietary": [3, 3, 4],
    "payment": [2, 3, 3], "location": [3, 3, 4], "praise": [5, 5, 5, 4],
}

ITEMS = ["Plain", "Vanilla", "Gluten-free", "Box of six"]
TERM_START = date(2026, 1, 26)  # 11 teaching weeks


def build_rows():
    rows = []
    comments = []
    for theme, texts in THEMES.items():
        for t in texts:
            comments.append((theme, t))
    rng.shuffle(comments)
    n_blank = 120 - len(comments)
    entries = comments + [(None, "")] * n_blank
    rng.shuffle(entries)

    for i, (theme, comment) in enumerate(entries, start=1):
        # weeks 1-11; queue/soldout complaints pile into weeks 5-8
        if theme in ("queue", "queue_price", "queue_soldout", "soldout") and rng.random() < 0.7:
            week = rng.randint(5, 8)
        else:
            week = rng.randint(1, 11)
        day = TERM_START + timedelta(weeks=week - 1, days=rng.randint(0, 4))
        slot = rng.choices(["Morning", "Lunch", "Afternoon"], weights=[2, 5, 3])[0]
        if theme in ("queue", "queue_price", "queue_soldout"):
            slot = rng.choices(["Lunch", "Afternoon"], weights=[8, 1])[0]
        if theme == "soldout":
            slot = "Afternoon" if rng.random() < 0.7 else "Lunch"
        item = rng.choices(ITEMS, weights=[5, 3, 2, 2])[0]
        if theme == "dietary" and rng.random() < 0.6:
            item = "Gluten-free"
        if theme is None:
            rating = rng.choices([3, 4, 5], weights=[2, 5, 4])[0]
        else:
            base = THEME_RATING[theme]
            rating = rng.choice(base)
            if 5 <= week <= 8 and theme in ("queue", "soldout") and rating > 2:
                rating -= 1
        spend = {"Plain": 2.5, "Vanilla": 2.5, "Gluten-free": 2.8, "Box of six": 12.0}[item]
        if rng.random() < 0.25:
            spend *= rng.choice([1, 2])
        rows.append({
            "response_id": f"R{i:03d}", "date": day.isoformat(), "time_slot": slot,
            "item": item, "rating": rating, "spend": round(spend, 2), "comment": comment,
        })

    # deliberate dirt: three missing ratings, one spend typo (blank comments exist already)
    for idx in (11, 47, 93):
        rows[idx]["rating"] = ""
    rows[66]["spend"] = 320.00
    return rows


ROWS = build_rows()
HEADERS = ["response_id", "date", "time_slot", "item", "rating", "spend", "comment"]

wb = Workbook()
ws = wb.active
ws.title = "Feedback"
ws.append(HEADERS)
for r in ROWS:
    ws.append([r[h] for h in HEADERS])
ws.freeze_panes = "A2"
wb.save(OUT / "Venture_Customer_Feedback.xlsx")

with open(OUT / "Venture_Customer_Feedback.csv", "w", newline="", encoding="utf-8") as f:
    w = csv.DictWriter(f, fieldnames=HEADERS)
    w.writeheader()
    w.writerows(ROWS)

# ======================================================= 4. THE BUSINESS PLAN
make_docx(OUT / "Venture_Business_Plan.docx", "Business Plan: Gravitas & Co. Traditional Welsh Cakes", [
    ("h", "1. Executive summary"),
    ("p", "Gravitas & Co. will bring authentic, freshly griddled Welsh cakes to the heart of "
          "campus. Our mission is to become the go-to campus destination for authentic Welsh "
          "baking within the first year. The venture is projected to generate £1,800 per week "
          "in profit from launch."),
    ("h", "2. Market opportunity"),
    ("p", "The campus hosts 12,000 students. Assuming a conservative 5 per cent daily "
          "conversion, the stall will sell 600 cakes per day. Research shows 78 per cent of "
          "students would pay a premium for local Welsh produce (Cardiff Student Survey, "
          "2024). The campus coffee shop does not sell fresh Welsh cakes, so we have no "
          "competition."),
    ("h", "3. Product and pricing"),
    ("p", "A focused menu: plain, vanilla and seasonal specials, sold at £2.50 each, with a "
          "box of six for £12.00. Every cake is made fresh on the stall in front of the "
          "customer, which is our theatre and our marketing."),
    ("h", "4. Operations"),
    ("p", "The stall operates with one traditional cast-iron griddle producing eight cakes "
          "every six minutes, trading over a four-hour daily window (10:00 to 14:00), five "
          "days a week, run by the founder as sole operator to keep costs minimal. Batter is "
          "prepared each morning; a single local supplier provides flour and butter."),
    ("h", "5. Financial projections"),
    ("p", "At 600 cakes per day and a unit price of £3.20, weekly revenue reaches £9,600, "
          "yielding the £1,800 weekly profit stated above after ingredients. Annual profit is "
          "therefore £93,600 (52 weeks). These figures demonstrate the venture is highly "
          "attractive to investors."),
    ("h", "6. Marketing and loyalty"),
    ("p", "Launch marketing relies on the irresistible smell and word of mouth. A loyalty "
          "scheme will collect student names and email addresses at the stall; to maximise "
          "value, the scheme shares the list with local partner businesses for cross-promotion."),
    ("h", "7. Risks"),
    ("p", "We do not anticipate significant risks. Welsh cakes are beloved, the campus is "
          "busy, and the product sells itself."),
])

# ============================================================== 5. THE RUBRIC
make_pdf(OUT / "Venture_Rubric.pdf", "Business Plan Rubric", [
    ("p", "Score each criterion 1 (absent or wrong) to 5 (strong, evidenced). Every score "
          "must cite the sentence or figure in the plan that justifies it."),
    ("t", [
        ["Criterion", "What a 5 looks like", "What a 1 looks like"],
        ["Market evidence",
         "Claims traced to verifiable sources; demand quantified honestly",
         "Invented or uncited statistics; wishful conversion rates"],
        ["Financial realism",
         "Revenue, costs and profit distinguished; prices consistent throughout",
         "Revenue presented as profit; numbers that contradict each other"],
        ["Operations and compliance",
         "Capacity arithmetic holds; food registration, hygiene and allergen "
         "duties (incl. Natasha's Law) addressed",
         "Production ceiling ignored; no mention of legal basics"],
        ["Competition and differentiation",
         "Competitors and substitutes named with a credible answer to each",
         "Competition denied or dismissed in a sentence"],
        ["Risk and contingency",
         "Single points of failure identified with mitigation; seasonality planned for",
         "“We do not anticipate significant risks”"],
    ]),
    ("p", "Total /25. Below 15: do not show an investor. Below 10: do not show your mother."),
])

# ==================================================== 6. THE FEEDBACK SUMMARY
make_pdf(OUT / "Venture_Feedback_Summary.pdf", "Customer Feedback Summary: One Term of Trading", [
    ("p", "Summary of 120 customer feedback responses (104 with comments) collected across an "
          "11-week teaching term at the stall. Average rating 3.2 out of 5, dipping below 2.8 "
          "in weeks 5 to 8 as queue and sell-out complaints peaked, recovering to about 4 by "
          "the end of term."),
    ("h", "Themes, ranked"),
    ("b", "Queues and waiting time: 34 comments (31%). Lunchtime waits of 15 to 25 minutes; "
          "customers report walking away and defecting to competitors."),
    ("b", "Price and value: 22 comments (20%). Split roughly evenly between “too dear for "
          "every day” and “fair for fresh”."),
    ("b", "Sold out by early afternoon: 19 comments (17%). Trays empty by 13:00 most days; "
          "gluten-free stock gone by late morning."),
    ("b", "Dietary and allergens: 14 comments (13%). Requests for gluten-free volume, vegan "
          "options and visible allergen labelling."),
    ("b", "Payment: 9 comments (8%). Card machine reliability and a cash-only day caused "
          "friction and lost sales."),
    ("b", "Location and hours: 7 comments (6%). Requests for earlier opening, Fridays, and "
          "better signage."),
    ("b", "Praise with nothing actionable: 5 comments (5%)."),
    ("h", "The one thing to fix"),
    ("p", "Queues and sell-outs are the same capacity problem seen from two ends; together "
          "they account for 48 per cent of all comments. Any intervention that smooths the "
          "lunch spike (pre-ordering, pickup slots, a live sold-out board) attacks both."),
])

# ==================================================== 6b. THE SALES LOG (m8)
# Planted Lean Six Sigma patterns, all deterministic:
#  - Plain cakes ~60% of revenue (Pareto vital few among products)
#  - Weeks 5-6: griddle temperature fault -> special-cause dip in sales and
#    a spike in "Griddle temperature fault" waste
#  - One clerical outlier: 2026-03-04 Plain units 250 (typo for 25)
#  - Waste Pareto: Burnt + Broken ~80% of waste cost
PRODUCTS = {  # name: (base daily units, unit price)
    "Plain": (58, 2.50),
    "Vanilla": (18, 2.50),
    "Gluten-free": (9, 2.80),
    "Box of six": (7, 12.00),
}
WEEKDAY_F = {0: 0.85, 1: 1.0, 2: 1.2, 3: 1.25, 4: 0.7}  # Mon..Fri
WASTE_REASONS = [  # reason, weight, units-lost range
    ("Burnt on griddle", 5, (4, 10)),
    ("Broken in handling", 4, (3, 8)),
    ("Over-proofed batch", 2, (2, 5)),
    ("Dropped at counter", 1, (1, 3)),
    ("Rain-damaged stock", 1, (1, 4)),
]
INGREDIENT_COST = 0.80  # per cake, for cost-of-poor-quality

srng = random.Random(808)
sales_rows, waste_rows = [], []
for week in range(1, 12):
    for dow in range(5):
        day = TERM_START + timedelta(weeks=week - 1, days=dow)
        week_f = 0.55 if week in (5, 6) else 1.0  # griddle fault fortnight
        for name, (base, price) in PRODUCTS.items():
            units = max(0, round(base * WEEKDAY_F[dow] * week_f * srng.uniform(0.82, 1.18)))
            if day == date(2026, 3, 4) and name == "Plain":
                units = 250  # the typo (should be 25)
            channel = "Pre-order" if srng.random() < 0.2 else "Stall"
            sales_rows.append({
                "date": day.isoformat(), "product": name, "units": units,
                "unit_price": price, "revenue": round(units * price, 2),
                "channel": channel,
            })
        # waste log: most days one entry, griddle fortnight adds fault entries
        if srng.random() < 0.75:
            name_r, _, rng_r = srng.choices(WASTE_REASONS, weights=[w for _, w, _ in WASTE_REASONS])[0]
            lost = srng.randint(*rng_r)
            waste_rows.append({"date": day.isoformat(), "reason": name_r,
                               "units_lost": lost, "cost": round(lost * INGREDIENT_COST, 2)})
        if week in (5, 6) and dow in (1, 2, 3):
            lost = srng.randint(12, 22)
            waste_rows.append({"date": day.isoformat(), "reason": "Griddle temperature fault",
                               "units_lost": lost, "cost": round(lost * INGREDIENT_COST, 2)})

wb2 = Workbook()
ws_s = wb2.active
ws_s.title = "Sales"
S_HEAD = ["date", "product", "units", "unit_price", "revenue", "channel"]
ws_s.append(S_HEAD)
for r in sales_rows:
    ws_s.append([r[h] for h in S_HEAD])
ws_s.freeze_panes = "A2"
ws_w = wb2.create_sheet("Waste log")
W_HEAD = ["date", "reason", "units_lost", "cost"]
ws_w.append(W_HEAD)
for r in waste_rows:
    ws_w.append([r[h] for h in W_HEAD])
ws_w.freeze_panes = "A2"
wb2.save(OUT / "Venture_Sales_Log.xlsx")

# ======================================================== 7. THE THEME LIST
(OUT / "Venture_Theme_List.txt").write_text(
    "Complaint themes from one term of customer feedback (ranked)\n"
    "============================================================\n\n"
    "1. Long queues at lunchtime and selling out by 1pm (48% of comments;\n"
    "   two faces of the same capacity problem)\n"
    "2. Price and value: fine for a treat, too dear for daily (20%)\n"
    "3. Dietary and allergens: more gluten-free, vegan option, visible labels (13%)\n"
    "4. Payment friction: card machine reliability, no contactless day (8%)\n"
    "5. Location and hours: open earlier, trade Fridays, better signage (6%)\n\n"
    + FOOTER + "\n",
    encoding="utf-8",
)

print("Generated artifacts in", OUT)
for f in sorted(OUT.glob("Venture_*")):
    print(f"  {f.name}  {f.stat().st_size:,} bytes")
